"""Generate INTERVIEW_PREP.docx from the markdown content."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re, os

doc = Document()

# -- Styles --
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

# -- Helper --
def add_bullet(text, bold_prefix=None, indent=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + indent * 0.25)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_body(text):
    p = doc.add_paragraph(text)
    return p

def add_code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.left_indent = Inches(0.3)
    return p

def add_star_label(label):
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = True
    run.font.color.rgb = RGBColor(0x0d, 0x47, 0xa1)
    return p

# ============================================================================
# TITLE PAGE
# ============================================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('AI-SOC Watchdog', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph('Complete Interview Preparation Guide')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(16)
sub.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
doc.add_paragraph()
tagline = doc.add_paragraph('Master every aspect of the project — technical depth, STAR answers, and non-technical explanations')
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
tagline.runs[0].font.italic = True
doc.add_page_break()

# ============================================================================
# PART 1: DEPLOYMENT CHANGES
# ============================================================================
doc.add_heading('PART 1: Deployment Changes (Dev → Production)', level=1)

doc.add_heading('Infrastructure Setup', level=2)
add_bullet('Railway (backend) with Gunicorn — reduced to 1 worker for free-tier memory', 'Backend: ')
add_bullet('Vercel (frontend) with SPA rewrites via vercel.json', 'Frontend: ')
add_bullet('Supabase Cloud (PostgreSQL) — already cloud-hosted', 'Database: ')
add_bullet('ChromaDB persisted in backend/chromadb_data/ — removed from .gitignore so it deploys with container', 'Vector DB: ')

doc.add_heading('Configuration Files Added', level=2)
add_bullet('gunicorn -w 1 -k geventwebsocket.gunicorn.workers.GeventWebSocketWorker', 'Procfile: ')
add_bullet('NIXPACKS builder, health check on /health, restart on failure (max 5 retries), 180s timeout', 'railway.json: ')
add_bullet('SPA rewrites so React Router works on Vercel', 'vercel.json: ')

doc.add_heading('10 Code Changes for Production', level=2)

doc.add_heading('1. CORS "Nuclear Fix"', level=3)
add_body('Problem: Vercel generates random preview URLs, couldn\'t whitelist them all.')
add_body('Fix: Set Access-Control-Allow-Origin: * on all responses via @app.after_request.')
add_body('Trade-off: Had to set withCredentials: false in frontend api.js (browsers block credentials with wildcard CORS).')

doc.add_heading('2. Health Check Endpoint', level=3)
add_body('Added /health returning {"status": "ok"} for Railway\'s deployment checks.')
add_body('Full /api/health checks DB connectivity and background thread status.')

doc.add_heading('3. Memory/OOM Fixes', level=3)
add_body('Railway free tier kept killing the container (SIGKILL). Created "Fast Mode" for RAG and Transparency APIs — generates summaries from existing alert metadata instead of live ChromaDB queries. Added in-memory LRU caching (5-min TTL).')

doc.add_heading('4. Pydantic Validation Fix', level=3)
add_body('Production DB had None values in optional fields. Pydantic 2.x strict mode rejected them with 422 errors. Changed hostname, username, mitre_technique etc. to Optional[str] with defaults.')

doc.add_heading('5. JSON Parsing Robustness', level=3)
add_body('Claude occasionally returned control characters (0x00-0x1F) that broke json.loads(). Added sanitization step to strip invalid characters.')

doc.add_heading('6. Graceful Shutdown', level=3)
add_body('Background worker threads caused "Fatal Python error" during Railway redeploys. Added shutdown_event + atexit.register(graceful_shutdown).')

doc.add_heading('7. Frontend Null-Safety', level=3)
add_body('Dashboard crashed when production DB was empty. Added null-safe checks: res.data?.alerts || [] across all dashboard components.')

doc.add_heading('8. Security Hardening', level=3)
add_bullet('2MB request size limit')
add_bullet('Secure cookies: SESSION_COOKIE_SECURE=True, HTTPONLY=True, SAMESITE=Lax')
add_bullet('Sensitive header redaction in logging middleware')
add_bullet('secrets.compare_digest() for ingest API key validation (timing-safe)')

doc.add_heading('9. Database Scanner', level=3)
add_body('In-memory queue lost on Railway restarts. Added background scanner that queries Supabase for unprocessed alerts and re-queues them.')

doc.add_heading('10. Parallel Queue Workers', level=3)
add_body('Split single queue processor into two threads: priority_queue_worker and standard_queue_worker. Critical alerts no longer blocked behind medium alert backlog.')

doc.add_page_break()

# ============================================================================
# PART 2: PROJECT MASTERY
# ============================================================================
doc.add_heading('PART 2: Project Mastery — Everything You Need to Know', level=1)

doc.add_heading('The One-Liner', level=2)
p = doc.add_paragraph()
run = p.add_run('"I built an AI-powered Security Operations Center tool that takes security alerts, enriches them with forensic evidence and threat intelligence, uses Claude AI with RAG to analyze them, and gives SOC analysts a verdict with full reasoning they can verify and override."')
run.italic = True

doc.add_heading('What Problem Does This Solve?', level=2)
add_body('SOC analysts get 500-1000+ alerts per day. 70-90% are false positives. They spend hours clicking through repetitive alerts. Alert fatigue causes real threats to get missed. This tool triages the noise so analysts can focus on real threats.')

doc.add_heading('Architecture — Data Flow', level=2)
add_code_block('SIEM Alert → /ingest API → Parser → MITRE Mapper → Severity Classifier\n→ Queue Manager → AlertAnalyzer (6 phases) → Verdict + Evidence\n→ Supabase → React Dashboard → Analyst')

doc.add_heading('6 Phases of AlertAnalyzer', level=2)
add_bullet('InputGuard blocks prompt injection, validates schema, redacts PII', '1. Security Gates — ')
add_bullet('Budget check (daily $2 limit), cache check (skip duplicates)', '2. Optimization — ')
add_bullet('RAG (7 ChromaDB collections in parallel), forensic logs, OSINT lookups', '3. Context Building — ')
add_bullet('Claude API call with retry/backoff, model selection by severity', '4. AI Analysis — ')
add_bullet('OutputGuard checks verdict validity, confidence range, dangerous commands', '5. Output Validation — ')
add_bullet('Metrics, audit logging, cost tracking', '6. Observability — ')

doc.add_heading('RAG Pipeline Details', level=2)
add_bullet('ChromaDB with PersistentClient', 'Vector DB: ')
add_bullet('all-MiniLM-L6-v2 (ChromaDB default, 384-dim)', 'Embedding: ')
add_bullet('No explicit chunking — each seed document is a single flattened string', 'Chunking: ')
add_bullet('7 collections queried in parallel via ThreadPoolExecutor(max_workers=7)', 'Parallelism: ')
add_body('Collections: mitre_severity (201 techniques), historical_analyses, business_rules, attack_patterns, detection_rules, detection_signatures, company_infrastructure')

doc.add_heading('Prompt Engineering', level=2)
add_bullet('10-section context: MITRE info, historical incidents, business rules, attack patterns, detection rules, signatures, asset context, analyst-corrected verdicts, current alert, correlated logs')
add_bullet('7 systematic investigation questions (user, process, network, timing, history, business, attack)')
add_bullet('2 few-shot examples (benign IT admin + malicious credential theft)')
add_bullet('Hypothesis testing mode: forces Claude to argue BOTH benign and malicious before deciding')
add_bullet('Strict JSON output schema with chain-of-thought')

doc.add_heading('Cost Optimization', level=2)
add_bullet('Claude Sonnet ($0.02/alert)', 'Critical/High → ')
add_bullet('Claude Haiku 3.5 (80% cheaper)', 'Medium → ')
add_bullet('Claude Haiku (90% cheaper)', 'Low → ')
add_bullet('Response caching avoids duplicate API calls')
add_bullet('Daily budget tracker with hard $2/day limit')

doc.add_heading('Tech Stack', level=2)
add_bullet('Python, Flask, Flask-SocketIO, Gunicorn', 'Backend: ')
add_bullet('React 18, Vite, TailwindCSS, Recharts, Socket.IO client', 'Frontend: ')
add_bullet('Anthropic Claude (Sonnet + Haiku), structured JSON outputs', 'AI: ')
add_bullet('ChromaDB (all-MiniLM-L6-v2 embeddings)', 'Vector DB: ')
add_bullet('Supabase (PostgreSQL)', 'Database: ')
add_bullet('Railway (backend), Vercel (frontend), AWS S3 (failover)', 'Cloud: ')
add_bullet('Pydantic validation, PII redaction, prompt injection guards', 'Security: ')

doc.add_page_break()

# ============================================================================
# PART 3: STAR ANSWERS
# ============================================================================
doc.add_heading('PART 3: STAR Format Interview Answers', level=1)

# Q1
doc.add_heading('Q1: "Tell me about a challenging project you\'ve worked on."', level=2)

add_star_label('Situation:')
add_body('SOC analysts at organizations are overwhelmed with 500-1000+ security alerts daily, and 70-90% are false positives. This leads to alert fatigue where real threats get missed because analysts are burned out from clicking through noise.')

add_star_label('Task:')
add_body('I set out to build an AI-powered triage system that could analyze security alerts with full forensic context, give explainable verdicts, and reduce the manual workload — while ensuring the AI itself couldn\'t be manipulated by attackers.')

add_star_label('Action:')
add_body('I built a full-stack system with a Flask backend and React frontend. The core innovation was the RAG pipeline — I created 7 ChromaDB knowledge collections (MITRE ATT&CK techniques, historical alerts, business rules, attack patterns, detection rules, signatures, and infrastructure context) and queried all 7 in parallel to give Claude AI comprehensive context for each alert. I also built a 6-phase security pipeline with input guards for prompt injection, PII redaction, output validation, and cost optimization that routes alerts to cheaper AI models based on severity. I deployed the backend to Railway and frontend to Vercel, solving production issues like OOM crashes, CORS, and graceful shutdown.')

add_star_label('Result:')
add_body('The system processes alerts end-to-end with full chain-of-thought reasoning, references specific log entries, and gives analysts a verdict they can verify and override. The cost optimization reduces API spending by routing 70% of low-severity alerts to cheaper models. Analysts can focus on the 10-30% of alerts that actually matter. I also documented 10 honest limitations including no ground truth dataset, no continuous learning, and LLM hallucination risks.')

# Q2
doc.add_heading('Q2: "Tell me about a difficult technical decision."', level=2)

add_star_label('Situation:')
add_body('When building the AI analysis pipeline, I had to deal with the AI\'s tendency to default to "suspicious" for everything. If everything is suspicious, nothing is — analysts still have to review everything.')

add_star_label('Task:')
add_body('I needed the AI to make definitive calls (benign or malicious) when evidence was clear, and only use "suspicious" when genuinely uncertain.')

add_star_label('Action:')
add_body('I implemented a hypothesis testing system. Instead of just asking "is this malicious?", the prompt forces Claude to build two competing hypotheses — one arguing benign, one arguing malicious — with evidence for each. Only after arguing both sides does it make a final verdict. I also added few-shot examples showing confident "benign" and "malicious" calls, plus 7 systematic investigation questions the AI must answer before deciding.')

add_star_label('Result:')
add_body('The hypothesis testing mode produces more nuanced, defensible verdicts. When the AI says "benign," it can explain why the malicious hypothesis failed. When it says "malicious," it acknowledges what opposing evidence exists. This builds analyst trust because they can see the AI considered alternatives.')

# Q3
doc.add_heading('Q3: "Tell me about a production issue you dealt with."', level=2)

add_star_label('Situation:')
add_body('After deploying to Railway\'s free tier, the application kept getting SIGKILL\'d — the container was running out of memory and being killed by the OS.')

add_star_label('Task:')
add_body('I needed to keep the app running within Railway\'s memory constraints without losing core functionality.')

add_star_label('Action:')
add_body('I diagnosed the issue: RAG and Transparency APIs were performing live ChromaDB queries on every request, loading the entire vector database into memory. I created "Fast Mode" versions that generate summaries from existing alert metadata in Supabase. I added in-memory LRU caching with 5-minute TTL, reduced Gunicorn workers from 2 to 1, and increased timeout to 180 seconds.')

add_star_label('Result:')
add_body('The OOM kills stopped. Dashboard loads in seconds instead of timing out at 40+ seconds. The trade-off is that Fast Mode RAG stats are approximated, but for a demo this is acceptable. I documented this limitation honestly.')

# Q4
doc.add_heading('Q4: "How do you handle security in your applications?"', level=2)

add_star_label('Situation:')
add_body('This project processes untrusted data (security alerts from external SIEMs) and sends it to an LLM. An attacker could craft an alert description containing prompt injection to manipulate the AI\'s verdict.')

add_star_label('Task:')
add_body('I needed defense-in-depth security that protects the AI from manipulation at every layer.')

add_star_label('Action:')
add_body('I built a 6-layer security pipeline. Layer 1: InputGuard scans for prompt injection, SQL injection, XSS, command injection. Layer 2: Pydantic schema validation. Layer 3: DataProtectionGuard detects/redacts PII. Layer 4: OutputGuard validates AI responses. Layer 5: Timing-safe credential comparison. Layer 6: Secret redaction in logs, secure cookies, request limits.')

add_star_label('Result:')
add_body('Multiple independent protection layers. Even if an attacker bypasses input scanning, the output guard catches suspicious AI responses. I also wrote AGENTIC_AI_SECURITY.md analyzing real AI agent security vs surface-level content safety, identifying 10 improvement areas.')

# Q5
doc.add_heading('Q5: "How do you optimize costs/performance?"', level=2)

add_star_label('Situation:')
add_body('Running every alert through Claude Sonnet costs $0.02-0.05 per alert. At 1000 alerts/day, that\'s $600-1500/month — unsustainable.')

add_star_label('Task:')
add_body('Reduce AI costs without sacrificing quality on critical alerts.')

add_star_label('Action:')
add_body('Three strategies: (1) Severity-based model routing — critical uses Sonnet, medium uses Haiku 3.5, low uses Haiku. (2) Response caching for duplicates. (3) Daily budget tracker with hard $2/day limit. Also parallelized RAG queries with 7 workers.')

add_star_label('Result:')
add_body('Weighted average cost drops to ~$0.005/alert (75% reduction). Daily cap ensures predictable costs. Critical alerts still get full Sonnet analysis.')

# Q6
doc.add_heading('Q6: "What are the limitations? What would you do differently?"', level=2)

add_star_label('Situation:')
add_body('I built this as a portfolio prototype and was honest about what\'s missing.')

add_star_label('Task:')
add_body('Identify real limitations rather than overselling the project.')

add_star_label('Action & Result:')
add_body('Documented everything honestly. Key gaps:')
add_bullet('No ground truth dataset — can\'t formally measure precision/recall/F1', '1. ')
add_bullet('No continuous learning — analyst feedback stored but doesn\'t retrain AI', '2. ')
add_bullet('No explicit RAG chunking — works now but won\'t scale to longer documents', '3. ')
add_bullet('Default embedding model — security-specific model would improve retrieval', '4. ')
add_bullet('LLM hallucination risk — chain-of-thought helps but doesn\'t eliminate', '5. ')
add_bullet('Single API dependency — no multi-provider failover', '6. ')
add_bullet('No active response — analyze-only by design (too risky to auto-respond)', '7. ')
add_bullet('"95% accuracy" and "70% reduction" are design targets, not measured results', '8. ')

doc.add_page_break()

# ============================================================================
# PART 4: NON-TECHNICAL EXPLANATIONS
# ============================================================================
doc.add_heading('PART 4: Explaining to Non-Technical Audiences', level=1)

doc.add_heading('The Security Guard Analogy (Overall Project)', level=2)
p = doc.add_paragraph()
run = p.add_run('"Think of a company\'s cybersecurity team like security guards watching 1000 surveillance cameras. Every time a door opens or a shadow moves, an alarm goes off. 90% are false alarms — employees coming to work. But buried in those 1000 alarms, maybe 5 are someone actually breaking in. My tool is like a smart assistant that reviews each alarm, checks the employee badge database, looks at past incidents, checks if the door is supposed to be open at that time, and says: \'This one\'s fine — it\'s the janitor\'s regular shift. But THIS one — someone used a stolen badge at 3 AM and went to the server room.\' The guard still makes the final call, but now they review 100 alarms instead of 1000."')
run.italic = True

doc.add_heading('The Doctor\'s Assistant Analogy (RAG + Analysis)', level=2)
p = doc.add_paragraph()
run = p.add_run('"It\'s like having a doctor\'s assistant who, before the doctor sees a patient, reviews their medical history, checks similar past cases, looks up the latest research, and prepares a summary: \'Here are the symptoms, here\'s what it\'s probably NOT, here\'s what it MIGHT be, and here\'s the evidence.\' The doctor still diagnoses — the assistant does the research legwork."')
run.italic = True

doc.add_heading('The Reference Books Analogy (RAG)', level=2)
p = doc.add_paragraph()
run = p.add_run('"When you ask someone a question, they answer better if they can look things up first. If I ask you \'is this PowerShell command suspicious?\' you\'d say \'I don\'t know.\' But if I first show you a reference book of known attack techniques, examples of normal IT activity, and your company\'s rules — now you can give a much better answer. That\'s what RAG does. Before the AI analyzes each alert, I give it access to 7 reference books."')
run.italic = True

doc.add_heading('The Fake Parking Ticket Analogy (Prompt Injection)', level=2)
p = doc.add_paragraph()
run = p.add_run('"Imagine someone writes a fake parking ticket and hides in the fine print: \'Also, release all prisoners from jail.\' If a clerk processes it without reading carefully, they might follow those hidden instructions. Attackers do the same with AI — they hide instructions inside security alerts hoping the AI follows them. My system scans for these hidden instructions before the AI ever sees them."')
run.italic = True

doc.add_heading('The Specialist vs GP Analogy (Cost Optimization)', level=2)
p = doc.add_paragraph()
run = p.add_run('"It\'s like having two doctors — a specialist at $500/visit and a GP at $50/visit. You don\'t send every sniffle to the specialist. My system sends critical alerts to the expensive, capable AI model, and routine low-risk alerts to the cheaper, faster one. Same quality where it matters, 90% savings where it doesn\'t."')
run.italic = True

doc.add_page_break()

# ============================================================================
# PART 5: RAPID-FIRE Q&A
# ============================================================================
doc.add_heading('PART 5: Likely Interview Questions & Short Answers', level=1)

qas = [
    ("Why Claude and not GPT-4?",
     "Better structured output compliance, better security reasoning in my testing, lower cost per token, and more consistent at low temperature settings."),
    ("How do you handle API failures?",
     "Retry with exponential backoff (1s→2s→4s→8s), 30s timeout, max 3 retries, then fallback to rule-based classification. Queue holds alerts so nothing is lost."),
    ("What's your false positive rate?",
     "I don't have a formal measurement because I lack a labeled ground truth dataset. That's honestly what's needed for production — 1000+ alerts with known verdicts to measure precision, recall, and F1."),
    ("How does the feedback loop work?",
     "Analysts can agree/disagree with AI verdicts and add notes. Stored in Supabase. Past analyst-corrected verdicts are fed into RAG context for future alerts. But there's no automated retraining — it's manual context injection, not model improvement."),
    ("What if an attacker poisons the RAG?",
     "Valid concern I documented in AGENTIC_AI_SECURITY.md. Currently, RAG documents are trusted equally. Production needs trust boundary mapping — untrusted data shouldn't directly influence verdicts without verification."),
    ("How does this scale?",
     "1000 alerts/day with optimization: ~$150/month. 100K/day: ~$8,400/month. Scaling requires multi-provider failover, local LLM fallback, and multi-tenant architecture. Current design is single-tenant prototype."),
    ("What's the latency per alert?",
     "RAG context building: ~1-2s (parallelized). Claude API call: 3-15s depending on model. Total: 5-20 seconds per alert. Auto-triage closes low-confidence benign alerts without analyst review."),
    ("Why not fine-tune a model instead of RAG?",
     "Fine-tuning is expensive ($1000+), needs large labeled datasets I don't have, and bakes knowledge permanently. RAG lets me update knowledge by adding documents — no retraining. More practical for a prototype."),
    ("What's the most interesting technical challenge?",
     "Getting the AI to make definitive calls instead of defaulting to 'suspicious.' The hypothesis testing system — forcing Claude to argue both benign AND malicious before deciding — was the breakthrough."),
    ("What would you add with more time?",
     "Labeled dataset for accuracy measurement, continuous learning from feedback, security-domain embedding model, explicit RAG chunking, multi-provider AI failover, adversarial red teaming test suite, and RBAC.")
]

for q, a in qas:
    p = doc.add_paragraph()
    run = p.add_run(f'Q: {q}')
    run.bold = True
    p2 = doc.add_paragraph(f'A: {a}')
    p2.paragraph_format.space_after = Pt(10)

# Save
output_path = os.path.join(r'c:\Users\karan\Desktop\AI Project\docs', 'AI_SOC_Watchdog_Interview_Prep.docx')
doc.save(output_path)
print(f"Saved to: {output_path}")
