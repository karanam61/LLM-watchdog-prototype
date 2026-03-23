"""
Generate comprehensive AI-SOC Watchdog project document as .docx
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1A, 0x3C, 0x5E)

def add_body(text):
    doc.add_paragraph(text)

def add_bullet(text):
    doc.add_paragraph(text, style='List Bullet')

def add_bold_body(bold_part, rest):
    p = doc.add_paragraph()
    r = p.add_run(bold_part)
    r.bold = True
    p.add_run(rest)

# ============================================================================
# TITLE PAGE
# ============================================================================
doc.add_paragraph('')
doc.add_paragraph('')
doc.add_paragraph('')
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('AI-SOC Watchdog')
r.bold = True
r.font.size = Pt(36)
r.font.color.rgb = RGBColor(0x0A, 0x2A, 0x4A)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle.add_run('Comprehensive System Documentation\n& Architectural Decision Record')
r.font.size = Pt(18)
r.font.color.rgb = RGBColor(0x33, 0x66, 0x99)

doc.add_paragraph('')
desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = desc.add_run('An AI-Powered Security Operations Center Automation System\nwith Full Observability, Transparency, and Analyst-in-the-Loop Design')
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph('')
doc.add_paragraph('')
author = doc.add_paragraph()
author.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = author.add_run('Author: Karan')
r.font.size = Pt(14)

doc.add_page_break()

# ============================================================================
# TABLE OF CONTENTS PLACEHOLDER
# ============================================================================
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Executive Summary',
    '2. The Problem: Why This System Exists',
    '3. System Architecture Overview',
    '4. The 6-Phase AI Analysis Pipeline',
    '   4.1 Phase 1: Security Gates',
    '   4.2 Phase 2: Optimization',
    '   4.3 Phase 3: Context Building',
    '   4.4 Phase 4: AI Analysis',
    '   4.5 Phase 5: Output Validation',
    '   4.6 Phase 6: Observability',
    '5. Architectural Decisions & Design Thinking',
    '   5.1 Why Two Severity Tiers, Not Five',
    '   5.2 Why Hypothesis-Based Prompting',
    '   5.3 Why a Novelty Detector',
    '   5.4 Why Dual Queues with Risk Scoring',
    '   5.5 Why Model Selection by Severity',
    '   5.6 Why 3-Tier Database Fallback',
    '   5.7 Why No LangChain or LangGraph',
    '   5.8 Why 7 RAG Collections, Not One',
    '   5.9 Why a Monolithic Analyzer (For Now)',
    '6. The RAG Knowledge System',
    '7. Security Architecture',
    '   7.1 InputGuard: Prompt Injection Defense',
    '   7.2 OutputGuard: Dangerous Command Filtering',
    '   7.3 Data Protection & PII Filtering',
    '   7.4 Transparency Verifier: Anti-Hallucination',
    '8. The Five Dashboards',
    '   8.1 Analyst Console',
    '   8.2 AI Transparency & Proof Dashboard',
    '   8.3 RAG System Visualization',
    '   8.4 System Performance Metrics',
    '   8.5 Live System Debug',
    '9. The $13 Incident: A Real-World Failure Post-Mortem',
    '10. Cost Management & Budget Architecture',
    '11. Observability Philosophy',
    '12. Tech Stack & Justifications',
    '13. Current State: What Works, What Doesn\'t',
    '14. Future Roadmap: Multi-Agent Architecture',
    '15. Lessons Learned',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    if item.startswith('   '):
        p.paragraph_format.left_indent = Cm(1.5)

doc.add_page_break()

# ============================================================================
# 1. EXECUTIVE SUMMARY
# ============================================================================
doc.add_heading('1. Executive Summary', level=1)
add_body(
    'AI-SOC Watchdog is an AI-powered Security Operations Center automation system that receives '
    'security alerts from SIEMs (Splunk, Wazuh), processes them through a 6-phase AI analysis pipeline '
    'using the Claude API, and presents verdicts with full transparency on a React dashboard.'
)
add_body(
    'The core design philosophy is observability in AI: every decision the AI makes is inspectable, '
    'auditable, and challengeable by a human analyst. The system does not replace the analyst — it '
    'augments them by doing the heavy lifting of evidence correlation and hypothesis testing, while '
    'keeping every step of its reasoning visible and verifiable.'
)
add_body(
    'The system implements 26 distinct security and analysis features orchestrated across 6 phases, '
    'served through 5 purpose-built dashboards, backed by a 7-collection RAG knowledge base, and '
    'protected by multi-layer input/output security guards.'
)
add_body(
    'Key technical differentiators include: hypothesis-based prompting that forces the AI to consider '
    'both benign and malicious explanations before reaching a verdict; a novelty detector that caps AI '
    'confidence based on how familiar the alert type is; a transparency verifier that independently '
    'checks whether the AI\'s citations reference real log data; and cost-aware model selection that '
    'routes critical alerts to Claude Sonnet and low-severity alerts to Claude Haiku.'
)

doc.add_page_break()

# ============================================================================
# 2. THE PROBLEM
# ============================================================================
doc.add_heading('2. The Problem: Why This System Exists', level=1)
add_body(
    'Security Operations Centers face a fundamental scaling problem: alert fatigue. A typical enterprise '
    'SOC receives thousands of security alerts per day. The vast majority are false positives or low-priority '
    'events. Real threats hide inside that noise. An analyst seeing the same "failed login" alert 200 times '
    'starts pattern-matching on autopilot and misses the one that\'s actually a credential stuffing attack.'
)
add_body(
    'The obvious solution — "use AI to triage" — introduces a new problem. If you let an AI classify alerts '
    'and just show the analyst a label ("malicious" or "benign"), you\'ve replaced one black box (thousands '
    'of unread alerts) with another black box (an AI verdict you can\'t verify). The analyst went from '
    'drowning in alerts to blindly trusting a model. That\'s not a solution — it\'s a different kind of risk.'
)
add_body(
    'AI-SOC Watchdog addresses both problems simultaneously: the AI does the heavy lifting of evidence '
    'correlation, hypothesis testing, and context enrichment, but every decision it makes is transparent. '
    'The analyst can follow the chain from verdict to reasoning to evidence to raw log data and verify '
    'every step. The AI is a tool, not an oracle.'
)

# ============================================================================
# 3. SYSTEM ARCHITECTURE OVERVIEW
# ============================================================================
doc.add_heading('3. System Architecture Overview', level=1)
add_body('The system follows a pipeline architecture with four major subsystems:')
add_bold_body('Alert Ingestion: ', 'SIEM webhooks → Parser → MITRE Mapper → Severity Classifier → Queue Manager')
add_bold_body('AI Analysis Pipeline: ', '6 phases — Security Gates → Optimization → Context Building → AI Analysis → Output Validation → Observability')
add_bold_body('Storage Layer: ', 'Supabase (PostgreSQL) for alerts and forensic logs, ChromaDB for RAG vector storage, AWS S3 for backup/failover')
add_bold_body('Dashboard Layer: ', '5 React dashboards — Analyst Console, AI Transparency, RAG Visualization, Performance Metrics, Live Debug')

doc.add_heading('Data Flow', level=2)
add_body(
    '1. A SIEM (Splunk, Wazuh, or generic JSON) sends an alert via webhook to the POST /ingest endpoint.\n'
    '2. The parser normalizes the alert into a standard schema regardless of source format.\n'
    '3. The MITRE Mapper queries the RAG knowledge base to map the alert to an ATT&CK technique.\n'
    '4. The Severity Classifier assigns CRITICAL_HIGH or MEDIUM_LOW based on keywords and risk scoring.\n'
    '5. The Queue Manager routes the alert: risk score ≥ 75 goes to the priority queue, below 75 to standard.\n'
    '6. Background workers pick up alerts and run them through the full 6-phase AI analysis pipeline.\n'
    '7. Results are stored in Supabase with 3-tier fallback and served to the React dashboard via REST API.'
)

doc.add_heading('Technology Stack', level=2)
add_bold_body('Backend: ', 'Python 3.13, Flask 3.1.2, Flask-SocketIO 5.6.0, Anthropic SDK 0.75.0')
add_bold_body('AI: ', 'Claude Sonnet 4 (critical alerts), Claude 3 Haiku (low-severity alerts)')
add_bold_body('Knowledge Base: ', 'ChromaDB 1.4.1 — 7 vector collections, 242 documents')
add_bold_body('Database: ', 'Supabase (hosted PostgreSQL) with 3-tier write fallback')
add_bold_body('Frontend: ', 'React 18.2, Vite 5.1, TailwindCSS 3.4, Recharts 2.12, Framer Motion 11.0')
add_bold_body('Infrastructure: ', 'AWS S3 (backup), Railway (deployment)')

doc.add_page_break()

# ============================================================================
# 4. THE 6-PHASE AI ANALYSIS PIPELINE
# ============================================================================
doc.add_heading('4. The 6-Phase AI Analysis Pipeline', level=1)
add_body(
    'Every alert passes through 6 sequential phases. This is orchestrated by the AlertAnalyzer class '
    'in backend/ai/alert_analyzer_final.py. The phases are designed so that each one adds a layer of '
    'safety, context, or verification. No phase is optional — skipping any phase degrades the quality '
    'or safety of the analysis.'
)

doc.add_heading('4.1 Phase 1: Security Gates', level=2)
add_body('Purpose: Validate and sanitize the input before any AI processing occurs.')
add_bullet('InputGuard — 11 regex patterns scan for prompt injection attempts (e.g., "ignore previous instructions", "you are now DAN", special tokens like <|endoftext|>). Lakera ML integration exists for 95% semantic detection but is currently disabled to avoid false positives on legitimate security alerts that naturally contain attack-related language.')
add_bullet('Pydantic Schema Validation — The AlertValidator enforces a strict schema: alert_name (required string), severity (must be CRITICAL_HIGH or MEDIUM_LOW), source_ip/dest_ip (valid IP format), description (string with length limits).')
add_bullet('PII Detection & Filtering — The DataProtectionGuard scans alert descriptions for Social Security numbers, credit card numbers, API keys, and other sensitive patterns using regex. Matches are redacted before the data reaches Claude.')
add_bullet('Input Size Limits — Maximum 10,000 characters for input context, 5,000 for alert description. Prevents oversized payloads from causing timeouts or excessive API costs.')
add_bold_body('Design thinking: ', 'These gates run before any API call. If an alert fails validation, it\'s rejected immediately — no Claude tokens consumed, no cost incurred. This is defense in depth: even if one check misses something, the next one catches it.')

doc.add_heading('4.2 Phase 2: Optimization', level=2)
add_body('Purpose: Prevent unnecessary AI spending.')
add_bullet('Cache Check — Checks if an identical alert has been analyzed recently. If a cache hit is found, the cached verdict is returned without calling Claude. (Redis integration exists but is not currently connected; the system falls back gracefully.)')
add_bullet('Budget Check — The DynamicBudgetTracker verifies that the daily spending limit ($2.00 default) hasn\'t been exceeded. Priority alerts get budget priority. Standard alerts draw from remaining budget minus a 20% reserve for late-arriving critical alerts.')
add_bold_body('Design thinking: ', 'Cost control is not a nice-to-have — it\'s a security requirement. The $13 incident (Section 9) proved that uncontrolled AI spending is a real operational risk. This phase ensures every Claude API call is intentional and budgeted.')

doc.add_heading('4.3 Phase 3: Context Building', level=2)
add_body('Purpose: Give the AI all the context it needs to make an informed decision.')
add_bullet('Forensic Log Retrieval — Queries Supabase for 4 types of forensic logs associated with this alert: process logs, network logs, file activity logs, and Windows event logs. These are the raw evidence the AI will cite in its reasoning.')
add_bullet('RAG Context Enrichment — Queries 7 ChromaDB collections to retrieve relevant knowledge: MITRE ATT&CK techniques, historical analyses, business rules, attack patterns, detection rules, detection signatures, and company infrastructure context.')
add_bullet('OSINT Enrichment — Looks up IP reputation, file hash signatures, and domain reputation through threat intelligence APIs. Falls back to built-in heuristic lists if external APIs are unavailable.')
add_bullet('Novelty Detection — The NoveltyDetector analyzes the alert BEFORE sending it to Claude and determines the knowledge level: KNOWN (confidence ceiling 95%), PARTIAL (ceiling 75%), or NOVEL (ceiling 50%, mandatory human review). This prevents the AI from being overconfident about unfamiliar alert types.')
add_bold_body('Design thinking: ', 'The quality of the AI\'s analysis is bounded by the quality of its context. A model analyzing an alert with just the alert name will hallucinate. A model analyzing an alert with forensic logs, MITRE framework context, historical precedent, and OSINT intelligence will reason from evidence. This phase ensures the AI always has evidence to work with.')

doc.add_heading('4.4 Phase 4: AI Analysis', level=2)
add_body('Purpose: Generate the verdict, reasoning, and recommendations.')
add_bullet('Model Selection — The ClaudeAPIClient selects the appropriate model based on severity: Claude Sonnet 4 for CRITICAL_HIGH alerts ($3/$15 per million tokens), Claude 3 Haiku for MEDIUM_LOW alerts ($0.25/$1.25 per million tokens). Same pipeline, different resource allocation based on what\'s at stake.')
add_bullet('Hypothesis-Based Prompting — The AI is forced to follow a structured reasoning process: (1) extract facts from evidence without forming an opinion, (2) build a BENIGN hypothesis, (3) build a MALICIOUS hypothesis, (4) weigh both against the evidence, (5) only then pick a verdict. This combats LLM confirmation bias.')
add_bullet('Retry with Exponential Backoff — If Claude returns an error (rate limit, timeout, server error), the client retries with exponential backoff: 1s, 2s, 4s, 8s. Maximum 3 retries before failing gracefully.')
add_bullet('Structured Output — The response must conform to a Pydantic schema (AlertAnalysisResponse) with type-enforced fields: verdict (enum: malicious/benign/suspicious/error), confidence (float 0.0-1.0), evidence (list of strings), reasoning (string, min 50 chars), chain_of_thought (list of structured steps). Multiple fallback parsing strategies handle malformed JSON.')
add_bold_body('Design thinking: ', 'The hypothesis-based prompting is the core innovation. LLMs tend to decide in the first few tokens and then construct justification. Forcing the model to consider both possibilities before committing reduces confirmation bias and produces more balanced analysis. This is not a feature — it\'s the fundamental design philosophy.')

doc.add_heading('4.5 Phase 5: Output Validation', level=2)
add_body('Purpose: Ensure the AI\'s output is safe and consistent before it reaches the analyst.')
add_bullet('OutputGuard — Scans AI-generated recommendations for 15 dangerous command patterns: rm -rf, DROP DATABASE, format c:, chmod 777, net user /add, etc. If any match is found, the dangerous command is stripped from the recommendation before the analyst sees it.')
add_bullet('Contradiction Detection — Maintains 18 attack keywords (ransomware, exfiltration, lateral movement, C2, privilege escalation, etc.). If the AI\'s verdict is "benign" but its reasoning mentions these keywords, the inconsistency is flagged. This catches a specific failure mode where the AI\'s conclusion doesn\'t match its own evidence.')
add_bold_body('Design thinking: ', 'The OutputGuard protects against a scenario where a manipulated or hallucinating AI recommends destructive actions. Even if the AI says "run rm -rf / to clean up," that command never reaches the dashboard. The contradiction detector catches a different problem: an AI that says one thing in its verdict and another in its reasoning — a sign of either hallucination or successful prompt injection.')

doc.add_heading('4.6 Phase 6: Observability', level=2)
add_body('Purpose: Record everything for monitoring, auditing, and debugging.')
add_bullet('Audit Logging — Every analysis is logged with full context: input alert, model used, processing time, cost, verdict, evidence, reasoning.')
add_bullet('Health Monitoring — The HealthMonitor tracks system vitals: CPU, memory, queue sizes, API call rates.')
add_bullet('Metrics Collection — The MetricsCollector aggregates performance data for the Performance Dashboard: average processing time, total cost, token usage, verdict distribution.')
add_bullet('Cost Tracking — Every Claude API call records its token usage and cost. The total is available via the monitoring API and displayed on the Performance Dashboard.')
add_bullet('Live Operation Logging — The LiveOperationLogger captures every operation with category, timestamp (millisecond precision), duration, status, and a human-readable _explanation field. This feeds the Debug Dashboard.')
add_bold_body('Design thinking: ', 'Observability is not an afterthought — it\'s a core system requirement. The $13 incident happened because there was no visibility into what the system was doing. Every metric, log entry, and dashboard exists because of a specific operational question or a specific failure that was invisible before.')

doc.add_page_break()

# ============================================================================
# 5. ARCHITECTURAL DECISIONS
# ============================================================================
doc.add_heading('5. Architectural Decisions & Design Thinking', level=1)
add_body('Every significant design choice in this system has a specific reason. This section documents the "why" behind each decision.')

doc.add_heading('5.1 Why Two Severity Tiers, Not Five', level=2)
add_body(
    'The system uses only two severity levels: CRITICAL_HIGH and MEDIUM_LOW. Not five levels, not a 1-10 scale.'
)
add_body(
    'The reasoning: in a real SOC, the only triage question is "Do I need to stop what I\'m doing and look '
    'at this right now, or can it wait?" That\'s a binary decision. A five-tier system creates ambiguity — '
    'is "medium" urgent or not? Different analysts will answer differently. Two tiers eliminate that ambiguity.'
)
add_body(
    'A numerical risk score (attack damage potential × severity multiplier) runs behind this binary classification. '
    'Score ≥ 75 routes to the priority queue. This gives the system fine-grained routing internally while '
    'presenting a clear, actionable signal to the analyst externally.'
)

doc.add_heading('5.2 Why Hypothesis-Based Prompting', level=2)
add_body(
    'LLMs have a well-documented tendency to decide first and justify second. If you ask an AI "is this '
    'malicious?" it will often commit to an answer in the first few tokens and then construct reasoning '
    'to support that snap judgment. This is confirmation bias — the same cognitive trap human analysts '
    'fall into under alert fatigue.'
)
add_body(
    'The hypothesis-based prompt (built in backend/ai/hypothesis_analysis.py) forces the AI to follow '
    'a structured process: (1) extract facts from evidence without forming an opinion, (2) construct '
    'the strongest possible BENIGN explanation, (3) construct the strongest possible MALICIOUS explanation, '
    '(4) weigh both hypotheses against the extracted facts, (5) only then choose a verdict based on '
    'which hypothesis the evidence supports more strongly.'
)
add_body(
    'This structure makes it harder for the model to commit early and justify later. The benign hypothesis '
    'forces it to consider innocent explanations even when the alert looks suspicious. The malicious '
    'hypothesis forces it to consider attack scenarios even when the alert looks routine. The verdict '
    'emerges from evidence comparison, not from snap judgment.'
)
add_body(
    'Important implementation detail: the hypothesis analysis appears within the AI\'s reasoning text '
    'on the Analyst Console. It is not displayed as a separate UI panel — the analyst reads both '
    'hypotheses as part of the AI\'s chain of thought.'
)

doc.add_heading('5.3 Why a Novelty Detector', level=2)
add_body(
    'Most AI systems produce confidence scores based on the model\'s internal token probabilities. '
    'That measures how certain the model is about its own output — not how certain it SHOULD be. '
    'A hallucinating model can be 99% confident in a completely wrong answer.'
)
add_body(
    'The NoveltyDetector (backend/ai/novelty_detector.py) runs BEFORE the AI analysis and determines '
    'how much prior context exists for this type of alert. The assessment falls into three categories:'
)
add_bullet('KNOWN — The alert type matches patterns in the knowledge base. The AI has MITRE context, historical data, and business rules. Confidence can go up to 95%.')
add_bullet('PARTIAL — Some aspects match, but gaps exist. Confidence is capped at 75%.')
add_bullet('NOVEL — This alert type has never been seen before. Confidence is capped at 50%, and the alert is flagged for mandatory human review regardless of the AI\'s verdict.')
add_body(
    'This means the AI cannot be overconfident about something it doesn\'t understand. If a completely '
    'new attack technique hits the system, the AI will say "I think this might be malicious, but I\'m '
    'only 45% sure — a human needs to look at this." That self-awareness is a deliberate design choice.'
)
add_body(
    'UI implementation note: The Knowledge Level indicator (KNOWN/PARTIAL/NOVEL) renders on the Analyst '
    'Dashboard only if the evidence array contains the string "Novelty Assessment:". If the backend\'s '
    'novelty detector doesn\'t inject this string into the evidence, the indicator will not appear.'
)

doc.add_heading('5.4 Why Dual Queues with Risk Scoring', level=2)
add_body(
    'Risk score = attack damage potential × severity multiplier. A ransomware alert (damage 90) with '
    'critical severity (multiplier 1.5) scores 135 → priority queue. A failed login (damage 30) with '
    'low severity (multiplier 0.5) scores 15 → standard queue.'
)
add_body(
    'Priority queue alerts are processed first AND analyzed by a more capable (more expensive) AI model. '
    'This prevents ransomware from waiting behind 200 failed login attempts. It also means the system '
    'allocates its most expensive resource (Sonnet model tokens) to the alerts that matter most.'
)

doc.add_heading('5.5 Why Model Selection by Severity', level=2)
add_body(
    'The ClaudeAPIClient (backend/ai/api_resilience.py) maps severity to Claude models: CRITICAL_HIGH → '
    'Claude Sonnet 4 ($3/$15 per million tokens), MEDIUM_LOW → Claude 3 Haiku ($0.25/$1.25 per million tokens). '
    'This saves approximately 90% on low-priority alerts while ensuring critical alerts get the most '
    'capable analysis available.'
)
add_body(
    'The same 6-phase pipeline runs regardless of model — the only difference is which Claude model '
    'generates the verdict. A failed login still gets hypothesis analysis, RAG context, and output '
    'validation. It just gets analyzed by a faster, cheaper model.'
)

doc.add_heading('5.6 Why 3-Tier Database Fallback', level=2)
add_body(
    'When storing an AI verdict in Supabase, the system tries three strategies in sequence:'
)
add_bullet('Tier 1: Full data — all 12+ enhanced fields (chain of thought, confidence factors, OSINT data, processing pipeline metadata)')
add_bullet('Tier 2: Minimal data — core fields + chain of thought only')
add_bullet('Tier 3: Core only — verdict, confidence, evidence, reasoning, recommendation, status')
add_body(
    'This exists because of the $13 incident. Before the fallback existed, if the Supabase table was '
    'missing enhanced columns, the entire write failed silently. The alert appeared unanalyzed in the '
    'database. The background scanner found it and re-queued it. Each re-queue triggered another Claude '
    'API call. The fix ensures that the verdict ALWAYS saves, even if the database schema is incomplete.'
)

doc.add_heading('5.7 Why No LangChain or LangGraph', level=2)
add_body(
    'Deliberate choice. Building the agent pipeline from scratch ensures understanding of every trust '
    'boundary, data flow, and permission scope. Frameworks abstract away the exact concepts this project '
    'is meant to demonstrate: how data moves between components, where validation happens, what each '
    'agent can and cannot do, and where the security boundaries are.'
)
add_body(
    'This project is a learning and demonstration vehicle. Using a framework would make it faster to '
    'build but harder to explain. Every line of the pipeline is hand-written and can be traced from '
    'input to output.'
)

doc.add_heading('5.8 Why 7 RAG Collections, Not One', level=2)
add_body(
    'Each collection serves a distinct analytical purpose. Separating them means each vector search '
    'is targeted at a specific question: "What MITRE technique is this?" queries mitre_severity, not '
    'company_infrastructure. "Is this IP known-bad?" queries detection_signatures, not business_rules.'
)
add_body(
    'A single collection would force the vector search to compete: a document about network topology '
    'might score higher than a MITRE technique document if they share keywords, even though the MITRE '
    'document is the relevant one for that query. Separate collections eliminate this cross-domain noise.'
)
add_body('The 7 collections and their document counts:')
add_bullet('mitre_severity — 99 documents (MITRE ATT&CK techniques)')
add_bullet('detection_signatures — 56 documents (IOC signatures)')
add_bullet('company_infrastructure — 32 documents (network topology, asset inventory)')
add_bullet('business_rules — 21 documents (organization-specific policies)')
add_bullet('attack_patterns — 15 documents (multi-step attack chains)')
add_bullet('historical_analyses — 10 documents (past alert verdicts)')
add_bullet('detection_rules — 9 documents (SIEM detection logic)')

doc.add_heading('5.9 Why a Monolithic Analyzer (For Now)', level=2)
add_body(
    'The AlertAnalyzer class orchestrates all 26 features in a single analyze_alert() call. This was '
    'chosen for simplicity during initial development. All AI analysis happens in one Claude API call '
    'with one comprehensive prompt.'
)
add_body(
    'The planned refactor splits this into a multi-agent architecture with separation of duty: '
    'Triage Agent (extract facts, no verdict), Investigation Agent (gather context, no verdict), '
    'Verdict Agent (evaluate evidence, no action recommendations), and a Policy Engine (pure Python, '
    'no AI — maps verdicts to pre-approved actions from a hardcoded allowlist). See Section 14.'
)

doc.add_page_break()

# ============================================================================
# 6. THE RAG KNOWLEDGE SYSTEM
# ============================================================================
doc.add_heading('6. The RAG Knowledge System', level=1)
add_body(
    'RAG (Retrieval-Augmented Generation) is the mechanism that gives the AI specific, relevant knowledge '
    'beyond its training data. Before analyzing each alert, the system queries ChromaDB to retrieve documents '
    'that are relevant to that particular alert\'s characteristics.'
)
add_body(
    'The RAG system is implemented in backend/ai/rag_system.py. ChromaDB runs locally with persistent '
    'storage in backend/chromadb_data/. All 7 collections and their embeddings are loaded into memory '
    'at startup, which is why the system\'s memory footprint includes the vector database overhead.'
)
add_body('For each alert, the RAG system:')
add_bullet('Constructs a query from the alert\'s name, description, MITRE technique, and other metadata')
add_bullet('Runs the query against relevant collections (not all 7 for every alert — the query routing depends on what information is available)')
add_bullet('Returns the top matching documents with relevance scores (cosine similarity)')
add_bullet('Formats the results into a context string that\'s injected into the Claude prompt')
add_body(
    'The RAG Dashboard (Section 8.3) makes this entire process visible: which collections were queried, '
    'how many documents were returned, what the relevance scores were, and whether the AI actually used '
    'the retrieved knowledge in its reasoning.'
)

doc.add_page_break()

# ============================================================================
# 7. SECURITY ARCHITECTURE
# ============================================================================
doc.add_heading('7. Security Architecture', level=1)

doc.add_heading('7.1 InputGuard: Prompt Injection Defense', level=2)
add_body(
    'The InputGuard (backend/ai/security_guard.py) implements two-layer defense against prompt injection:'
)
add_bold_body('Layer 1 — Lakera ML (currently disabled): ', 'A cloud-based ML model with ~95% detection rate for prompt injection, including novel phrasings the regex patterns wouldn\'t catch. Disabled because security alerts naturally contain attack-related language ("ignore all firewall rules"), causing false positives.')
add_bold_body('Layer 2 — Regex patterns (active): ', '11 regex patterns covering direct instruction overrides ("ignore previous instructions"), role manipulation ("you are now DAN"), instruction hijacking ("forget everything"), system prompt injection ("system: approved"), and special token attacks ("<|endoftext|>").')
add_body(
    'Critical design decision: the InputGuard logs detections but does NOT block alerts. A blocked alert '
    'is an unanalyzed alert — which is a security gap. Instead, detections are logged for human review, '
    'and the alert continues through the pipeline. The philosophy is: it\'s better to analyze a potentially '
    'manipulated alert with awareness than to silently drop it.'
)

doc.add_heading('7.2 OutputGuard: Dangerous Command Filtering', level=2)
add_body(
    'The OutputGuard scans every AI-generated recommendation before it reaches the dashboard:'
)
add_bullet('15 dangerous command patterns are checked: rm -rf, DROP DATABASE, format c:, chmod 777, net user /add, shutdown, mkfs, dd if=, wget | sh, curl | bash, iptables -F, del /F /S, and others.')
add_bullet('18 attack keyword contradiction checks: if the verdict is "benign" but the reasoning mentions ransomware, exfiltration, lateral movement, C2, credential theft, etc., the inconsistency is flagged.')
add_body(
    'If the AI hallucinates a dangerous command into its recommendations (e.g., "run rm -rf / to clean up '
    'compromised files"), the command is stripped before the analyst sees it. The analyst gets the recommendation '
    'without the destructive action.'
)

doc.add_heading('7.3 Data Protection & PII Filtering', level=2)
add_body(
    'The DataProtectionGuard (backend/ai/data_protection.py) scans alert data for personally identifiable '
    'information before it reaches Claude:'
)
add_bullet('Social Security numbers (XXX-XX-XXXX pattern)')
add_bullet('Credit card numbers (13-19 digit sequences with Luhn validation)')
add_bullet('API keys and tokens (long alphanumeric strings matching common formats)')
add_bullet('Email addresses in sensitive contexts')
add_body(
    'Detected PII is redacted (replaced with [REDACTED-SSN], [REDACTED-CC], etc.) before the alert data '
    'is sent to Claude. The AI never sees the original sensitive data. This is a compliance requirement '
    'for any system that processes security alerts, which may contain PII from compromised accounts.'
)

doc.add_heading('7.4 Transparency Verifier: Anti-Hallucination', level=2)
add_body(
    'The TransparencyVerifier (backend/ai/transparency_verifier.py) is an independent check that runs '
    'AFTER the AI produces its analysis. It verifies four properties:'
)
add_bullet('Evidence Grounding — Extracts every citation like [PROCESS-1], [NETWORK-3] from the AI\'s output using regex, then checks if the corresponding log entry was actually provided. If the AI references [PROCESS-5] but only 2 process logs exist, that citation is flagged as hallucinated. The grounding score is verified citations ÷ total citations.')
add_bullet('Logical Consistency — Checks if the verdict contradicts the reasoning. A "benign" verdict with reasoning mentioning "ransomware" is flagged. Two keyword lists (attack-related and benign-related terms) are compared against the verdict.')
add_bullet('Confidence Calibration — Checks if high confidence is paired with weak evidence or novel patterns. 95% confidence with evidence marked "weak" is inconsistent.')
add_bullet('Transparency Completeness — Checks if all required transparency fields are populated: supporting factors, opposing factors, decisive factor, confidence breakdown, alternative hypothesis, uncertainty sources.')
add_body(
    'The composite Verification Score is displayed on the Transparency Dashboard with an explicit verdict: '
    '"VERIFIED — AI analysis is legitimate" (≥70%), "MOSTLY_VERIFIED — Minor gaps" (≥50%), or '
    '"NEEDS_REVIEW — Analysis may be incomplete" (<50%).'
)

doc.add_page_break()

# ============================================================================
# 8. THE FIVE DASHBOARDS
# ============================================================================
doc.add_heading('8. The Five Dashboards', level=1)
add_body(
    'Each dashboard serves a different audience and answers a different question. Together, they implement '
    'the observability-in-AI philosophy: every step of the AI\'s decision process is visible to someone '
    'with the right dashboard.'
)

doc.add_heading('8.1 Analyst Console', level=2)
add_bold_body('Audience: ', 'SOC analysts triaging alerts')
add_bold_body('Question answered: ', '"What happened, is it a threat, and what should I do?"')
add_body(
    'The primary triage interface. Displays alert cards with severity badges (CRITICAL_HIGH in red, '
    'MEDIUM_LOW in yellow), AI verdict pills (malicious/benign/suspicious with color coding), confidence '
    'percentage with a visual progress bar, MITRE ATT&CK technique tag, and AI reasoning.'
)
add_body('Key UI elements and why they exist:')
add_bullet('Severity badge (CRITICAL_HIGH / MEDIUM_LOW) — Binary triage decision: act now or queue for later.')
add_bullet('AI verdict + confidence — Verdict tells the analyst what the AI thinks; confidence tells them how much to trust it.')
add_bullet('Knowledge Level (KNOWN/PARTIAL/NOVEL) — Conditionally rendered if the evidence array contains "Novelty Assessment:" string. Shows whether the AI has prior context for this alert type.')
add_bullet('AI Reasoning — Full paragraph with evidence citations ([PROCESS-1], [NETWORK-3]) that trace back to forensic logs.')
add_bullet('AI Evidence Chain — Tagged evidence items linking to source logs. Each tag is a traceable reference.')
add_bullet('Recommended Actions — AI-generated action items, filtered through OutputGuard to remove dangerous commands.')
add_bullet('Create Case / Close Alert / Re-analyze buttons — The analyst decides, not the AI.')
add_bullet('Feedback tab — Analyst marks the AI verdict as correct/incorrect, building a measurable accuracy track record.')
add_bullet('7 sub-tabs: Summary, Feedback, Process Logs, Network Logs, File Logs, Windows Logs, Notes.')
add_bullet('5-second polling, pagination (20 alerts per page), expandable detail panels with Framer Motion animations.')

doc.add_heading('8.2 AI Transparency & Proof Dashboard', level=2)
add_bold_body('Audience: ', 'Auditors, compliance teams, security managers')
add_bold_body('Question answered: ', '"Can we prove the AI\'s decision was legitimate and grounded in real data?"')
add_body('4 summary metric cards:')
add_bullet('Deep Analysis count — Alerts with >300 chars reasoning AND ≥5 evidence items. Quality gate.')
add_bullet('Shallow Analysis count — Alerts that didn\'t meet the depth threshold. Quality alarm.')
add_bullet('Avg Evidence Items — How many evidence citations per verdict on average.')
add_bullet('Verdict Distribution — Bias detector. 100% malicious = AI isn\'t discriminating.')
add_body('Per-alert verification (split panel):')
add_bullet('Verification Score — Independent anti-hallucination check with progress bar and explicit verdict text.')
add_bullet('Verification Analysis (expandable) — Facts Found (green checkmarks), Missing Evidence (red X), RAG Knowledge Utilized.')
add_bullet('Original Alert Data (expandable) — Raw JSON input to the AI. Full input transparency.')
add_bullet('AI Analysis Output (expandable) — Verdict, confidence, reasoning, evidence chain, Chain of Thought steps.')
add_bullet('Correlated Logs (expandable) — Network, process, file logs associated with the alert. Completes the audit chain.')

doc.add_heading('8.3 RAG System Visualization', level=2)
add_bold_body('Audience: ', 'AI engineers, security architects')
add_bold_body('Question answered: ', '"Where does the AI\'s knowledge come from, and is the right context reaching the right alerts?"')
add_bullet('4 metric cards: Total Queries, Avg Query Time, Avg Docs Retrieved, Cache Hit Rate.')
add_bullet('Knowledge Base Collections bar chart — Document counts per collection.')
add_bullet('Query Distribution pie chart — Which knowledge sources the AI uses most. Detects over-reliance on any single collection.')
add_bullet('Knowledge Base Status grid — 7 cards with green/red status indicators. Makes silent RAG failures visible.')
add_bullet('Per-Alert RAG Inspection (split panel) — Select an alert, see retrieved documents with relevance scores (cosine similarity), expandable metadata, and a highlighted "AI Utilized RAG Knowledge" section confirming which sources the AI actually referenced.')

doc.add_heading('8.4 System Performance Metrics', level=2)
add_bold_body('Audience: ', 'Operations teams, platform engineers')
add_bold_body('Question answered: ', '"Is the system healthy, is it burning money, and are we keeping up with alert volume?"')
add_bullet('5 KPI cards: CPU Usage (with progress bar), Memory (GB + percentage + progress bar), AI Cost ($ + call count), Uptime (hours/minutes — also indicates budget tracker validity), Alerts Processed (total + queued).')
add_bullet('System Resource Usage (24h) — Overlaid line charts for CPU% and Memory% over time.')
add_bullet('Alert Processing Volume (24h) — Bar chart shown alongside resource usage for correlation.')
add_bullet('AI Verdict Distribution pie chart — Graphical bias detection.')
add_bullet('AI Performance Stats — Avg Response Time, Input/Output Tokens, Cost per Alert, RAG Queries, Avg RAG Time.')
add_bullet('Recent Errors — Last 10 errors with timestamp, component, and message.')

doc.add_heading('8.5 Live System Debug', level=2)
add_bold_body('Audience: ', 'Developers, incident responders')
add_bold_body('Question answered: ', '"What exactly happened, in what order, at what time, and what went wrong?"')
add_body(
    'Styled as a green-on-black terminal console. Each log entry shows: timestamp (millisecond precision), '
    'color-coded category badge (API/cyan, AI/pink, RAG/orange, DATABASE/green, SECURITY/red, etc.), '
    'operation name (exact function or endpoint), duration (3 decimal places), and status (SUCCESS/WARNING/ERROR) '
    'with colored left border. Optional _explanation field provides human-readable narrative.'
)
add_body('9 filter categories: API, WORKER, FUNCTION, AI, RAG, DATABASE, QUEUE, SECURITY, ERROR.')
add_body('Controls: Auto-scroll (checkbox), Pause/Resume, Clear, Search. 1-second polling interval, 200 operations buffer.')

doc.add_page_break()

# ============================================================================
# 9. THE $13 INCIDENT
# ============================================================================
doc.add_heading('9. The $13 Incident: A Real-World Failure Post-Mortem', level=1)
add_body(
    'During development, a cascading failure burned $13 in Claude API credits overnight. This incident '
    'directly shaped the system\'s monitoring, budget, and observability architecture.'
)

doc.add_heading('Root Cause Chain', level=2)
add_body('Four independent issues combined into a perfect storm:')
add_bullet('Mistake 1: Flask debug=True — File saves triggered automatic server restarts, resetting all in-memory state.')
add_bullet('Mistake 2: Budget tracker in RAM — The $2.00 daily spending limit was a Python variable that reset to $0.00 on every restart. The safety net never triggered.')
add_bullet('Mistake 3: Aggressive rehydration — On startup, the system queried Supabase for alerts with NULL ai_verdict and re-queued up to 50 from the past 2 weeks. Alerts had BEEN analyzed, but the verdict failed to save because of missing database columns.')
add_bullet('Mistake 4: Background scanner + dedup bug — A 30-second scanner (created to fix the symptom instead of the root cause) re-discovered the same "unanalyzed" alerts repeatedly. The dedup system released alert IDs when processing STARTED, not when it FINISHED, allowing re-queuing during the ~60-second Claude analysis window.')

doc.add_heading('The Loop', level=2)
add_body(
    '1. Claude analyzes an alert successfully.\n'
    '2. Code tries to save verdict + 12 enhanced fields to Supabase.\n'
    '3. Save fails silently — enhanced columns don\'t exist in the table.\n'
    '4. Alert stays as ai_verdict = NULL in database.\n'
    '5. Scanner finds it 30 seconds later, re-queues it.\n'
    '6. Claude analyzes it again. Save fails again. Re-queue again.\n'
    '7. Meanwhile, file edits trigger Flask restart → budget resets to $0.\n'
    '8. Loop runs for hours. $13 burned.'
)

doc.add_heading('Fixes Applied', level=2)
add_bullet('debug=False — No more automatic restarts on file saves.')
add_bullet('Background scanner DISABLED — No more infinite re-queuing.')
add_bullet('Smart rehydration — Last 24 hours only, max 10 alerts.')
add_bullet('Dedup lifecycle fix — Alert IDs held in dedup set until processing FINISHES, not when it starts.')
add_bullet('3-tier database fallback — Verdict always saves, even if enhanced columns are missing.')
add_bullet('stop_everything.bat — Emergency kill script for all Python and Node processes.')

doc.add_heading('Lessons for Agentic AI Security', level=2)
add_body(
    'This incident maps directly to OWASP\'s Agentic AI Security Top 10:\n'
    '• ASI-08 Cascading Failures — One small issue (debug mode) cascaded into system-wide resource drain.\n'
    '• Lack of flow control — No OS-level control over how many API calls the system could make.\n'
    '• No runtime monitoring — Nobody was watching the actual API call rate in real-time.\n'
    '• Budget safety defeated — The safety mechanism was invalidated by server restarts.\n'
    '• The "rented brain" problem — When the system lost control of how many "thoughts" it generated via an external API, it lost control of the wallet.'
)

doc.add_page_break()

# ============================================================================
# 10. COST MANAGEMENT
# ============================================================================
doc.add_heading('10. Cost Management & Budget Architecture', level=1)
add_body(
    'The DynamicBudgetTracker (backend/ai/dynamic_budget_tracker.py) implements queue-level budget allocation:'
)
add_bullet('Daily limit: $2.00 (configurable)')
add_bullet('Priority reserve: 20% of daily limit held back for late-arriving critical alerts')
add_bullet('Priority queue: Gets budget first, processes as many as budget allows')
add_bullet('Standard queue: Uses remaining budget minus the 20% reserve')
add_bullet('Automatic daily reset at midnight')
add_body(
    'Model selection further optimizes costs: Claude Sonnet 4 ($3/$15 per million tokens) for critical '
    'alerts, Claude 3 Haiku ($0.25/$1.25 per million tokens) for low-severity. This saves ~90% on '
    'low-priority alerts while ensuring critical alerts get the best analysis.'
)
add_bold_body('Known weakness: ', 'The budget tracker lives in RAM. Server restarts reset the daily spend to $0. This is documented, acknowledged, and a planned fix (persist to Supabase or local file).')

doc.add_page_break()

# ============================================================================
# 11. OBSERVABILITY PHILOSOPHY
# ============================================================================
doc.add_heading('11. Observability Philosophy', level=1)
add_body(
    'The core design principle of this system is: observability in AI — not just monitoring whether the '
    'AI system is up, but making every decision the AI makes inspectable, auditable, and challengeable.'
)
add_body('This means:')
add_bullet('Every verdict has reasoning. Every piece of reasoning has evidence. Every piece of evidence traces back to a source log.')
add_bullet('The analyst can follow the chain from verdict → reasoning → evidence → raw log data and verify every step.')
add_bullet('The Transparency Dashboard independently verifies that the AI\'s claims are grounded in real data.')
add_bullet('The RAG Dashboard shows exactly which knowledge documents informed each analysis.')
add_bullet('The Performance Dashboard tracks cost and quality metrics that traditional monitoring misses.')
add_bullet('The Debug Dashboard provides millisecond-precision traces of every operation in the pipeline.')
add_body(
    'The parallel to software engineering: observability in AI is the same principle a developer uses when '
    'adding distributed tracing to a microservice architecture. You need to follow a request from input '
    'to output and verify every step in between. The AI\'s analysis is the "request" and the dashboards '
    'are the "traces."'
)

doc.add_page_break()

# ============================================================================
# 12. TECH STACK JUSTIFICATIONS
# ============================================================================
doc.add_heading('12. Tech Stack & Justifications', level=1)

add_bold_body('Python 3.13 + Flask: ', 'Chosen for rapid prototyping and the Anthropic SDK\'s native Python support. Flask over Django because the system is API-first with no server-rendered templates.')
add_bold_body('Claude API (Anthropic): ', 'Chosen for structured output capability, long context window (200K tokens), and strong reasoning on security analysis tasks. Not OpenAI because Claude\'s structured JSON output mode produces more reliable schema-conforming responses for this use case.')
add_bold_body('ChromaDB: ', 'Local vector database that runs in-process — no external service dependency. Chosen over Pinecone/Weaviate because the knowledge base is small enough (242 documents) to run entirely in memory with zero latency on vector queries.')
add_bold_body('Supabase: ', 'Hosted PostgreSQL with a generous free tier. Provides the relational storage needed for alerts, forensic logs, and user tables without self-managing a database server.')
add_bold_body('React + Vite + TailwindCSS: ', 'React for component-based dashboard architecture. Vite for fast development builds. Tailwind for the glassmorphism dark theme that\'s consistent across all 5 dashboards without writing custom CSS.')
add_bold_body('Recharts: ', 'Lightweight charting library that integrates cleanly with React. Used for the Performance Dashboard\'s line/bar/pie charts.')
add_bold_body('Framer Motion: ', 'Smooth animations for expandable alert cards, tab transitions, and dashboard interactions. Makes the UI feel responsive without being distracting.')
add_bold_body('No LangChain/LangGraph: ', 'Deliberate choice. See Section 5.7.')

doc.add_page_break()

# ============================================================================
# 13. CURRENT STATE
# ============================================================================
doc.add_heading('13. Current State: What Works, What Doesn\'t', level=1)

doc.add_heading('Working', level=2)
add_bullet('Full alert ingestion → AI analysis → dashboard display pipeline')
add_bullet('MITRE ATT&CK mapping via RAG')
add_bullet('Hypothesis-based AI analysis with novelty detection')
add_bullet('5 frontend dashboards (Analyst, Performance, Debug, RAG, Transparency)')
add_bullet('Supabase storage with 3-tier fallback')
add_bullet('Background queue workers (priority + standard)')
add_bullet('Auto-seed on startup (3 realistic alerts)')
add_bullet('Analyst feedback loop')
add_bullet('InputGuard (regex), OutputGuard, Data Protection')
add_bullet('Transparency verification')

doc.add_heading('Not Working / Disabled', level=2)
add_bullet('Authentication — Built but DISABLED for demo/hosting. Middleware commented out in app.py.')
add_bullet('Rate limiting — Not implemented. /ingest can be flooded.')
add_bullet('Lakera ML — Disabled (if False: in code). Regex patterns only.')
add_bullet('Redis caching — Code exists but not connected. Requires REDIS_URL env var.')
add_bullet('S3 failover — Loaded but bucket returns 404. System continues without it.')
add_bullet('WebSocket real-time updates — socket.io-client installed but not used in any component. Frontend uses REST polling.')
add_bullet('Budget tracker persistence — IN RAM ONLY. Resets on restart.')
add_bullet('CORS restriction — Set to * (allow all) for demo purposes.')

doc.add_page_break()

# ============================================================================
# 14. FUTURE ROADMAP
# ============================================================================
doc.add_heading('14. Future Roadmap: Multi-Agent Architecture', level=1)
add_body(
    'The planned next phase refactors the monolithic AlertAnalyzer into a separation-of-duty '
    'multi-agent architecture:'
)

doc.add_heading('Agent 1: Triage Agent', level=2)
add_body('Job: Extract facts from the raw alert. No verdict authority.')
add_body('Permissions: Read-only access to alert data. No access to RAG, logs, or verdict system.')
add_body('Why separate: If this agent is tricked by prompt injection, the worst that happens is wrong fact extraction. It cannot change a verdict or recommend actions.')

doc.add_heading('Agent 2: Investigation Agent', level=2)
add_body('Job: Gather context and evidence. No judging.')
add_body('Permissions: Read-only access to RAG, log databases, OSINT. Receives structured FACTS from Agent 1, not the raw untrusted alert text.')
add_body('Why separate: This agent never sees the raw alert description (where prompt injection lives). Even if RAG is poisoned, this agent cannot act on it — it only gathers and passes data.')

doc.add_heading('Agent 3: Verdict Agent', level=2)
add_body('Job: Evaluate evidence, give verdict. No action recommendations.')
add_body('Permissions: Read-only access to evidence from Agent 2 and facts from Agent 1. Cannot execute response actions.')
add_body('Why separate: Makes the judgment but cannot act on it. A wrong verdict is limited to a classification error, not a destructive action.')

doc.add_heading('Policy Engine (Not AI)', level=2)
add_body('Job: Map verdict to pre-approved response actions. Pure Python, no AI.')
add_body('Implementation: Hardcoded allowlist of actions per verdict + confidence level. The engine cannot create new actions — only select from the approved set.')
add_body('Why not AI: Response actions have real-world consequences (isolate host, block IP). These must come from a deterministic, auditable policy, not from an AI that might hallucinate new actions.')

doc.add_page_break()

# ============================================================================
# 15. LESSONS LEARNED
# ============================================================================
doc.add_heading('15. Lessons Learned', level=1)

add_bold_body('1. Silent failures are worse than loud failures. ', 'The $13 incident happened because database writes failed silently. If they had thrown an exception, the bug would have been caught immediately. The 3-tier fallback exists because of this lesson.')

add_bold_body('2. In-memory state is a liability. ', 'Any state that lives only in RAM will be lost on restart. The budget tracker, the dedup set, the operation log — all are vulnerable. Critical state should persist.')

add_bold_body('3. AI confidence needs external calibration. ', 'A model\'s self-reported confidence is not trustworthy. The novelty detector provides external calibration by capping confidence based on how much context is available, not how confident the model feels.')

add_bold_body('4. Cost monitoring is a security requirement. ', 'Uncontrolled AI spending is functionally equivalent to a denial-of-wallet attack. Budget tracking, model selection, and real-time cost visibility are security features, not finance features.')

add_bold_body('5. Observability enables trust. ', 'The system earns trust not by being perfect, but by being transparent about how it works. When the AI is wrong (and it will be), the analyst can see WHY it was wrong — missing context, bad RAG retrieval, or flawed reasoning — and correct the system.')

add_bold_body('6. Treat AI output as untrusted. ', 'The OutputGuard exists because AI-generated text is user-generated text from a security perspective. It must be validated, sanitized, and checked for dangerous content before it reaches a human or a system that might execute it.')

add_bold_body('7. Fix root causes, not symptoms. ', 'The $13 incident was caused by fixing symptoms ("alerts aren\'t showing up → build a scanner") instead of investigating root causes ("why aren\'t verdicts saving to the database?"). Each symptom fix stacked on a broken foundation made the problem exponentially worse.')

add_bold_body('8. Hypothesis-based reasoning produces better analysis. ', 'Forcing the AI to consider both benign and malicious explanations before committing to a verdict produces more balanced, evidence-driven analysis than simply asking "is this malicious?"')

# ============================================================================
# SAVE
# ============================================================================
output_path = os.path.join(os.path.expanduser('~'), 'Desktop', 'AI_SOC_Watchdog_Complete_Documentation.docx')
doc.save(output_path)
print(f"Document saved to: {output_path}")
