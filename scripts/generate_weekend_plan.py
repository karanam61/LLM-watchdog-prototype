"""Generate Weekend Study Plan .docx"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

def bold_body(bold_text, normal_text):
    p = doc.add_paragraph()
    r = p.add_run(bold_text)
    r.bold = True
    p.add_run(normal_text)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def red_text(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    return p

def green_text(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.color.rgb = RGBColor(0x00, 0x7A, 0x33)
    return p

def code(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.left_indent = Inches(0.3)
    return p

# ============================================================================
# TITLE
# ============================================================================
doc.add_paragraph()
title = doc.add_heading('Weekend Survival Plan', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph('Linux + Docker + AWS Deployment — 2 Hours/Day')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(14)
sub.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
doc.add_page_break()

# ============================================================================
# HONEST ASSESSMENT
# ============================================================================
doc.add_heading('Honest Assessment: Is This Possible?', level=1)

red_text('The raw material = 25 concepts × 15-30 min each = 6-12 hours hands-on.')
red_text('You have = 2 days × 2 hours = 4 hours total.')

doc.add_paragraph()
green_text('YES, it\'s doable — but only if you change the approach:')
doc.add_paragraph()

bullet('SKIP typing every command — read the concept, type only the 3-4 key commands, move on', 'Strategy 1: ')
bullet('GROUP related concepts — users + groups + permissions is ONE session, not three', 'Strategy 2: ')
bullet('INTERVIEW ANSWERS FIRST — if you can explain it, you can do it. Hands-on proves depth, but verbal fluency wins interviews', 'Strategy 3: ')
bullet('DOCKER > LINUX for this job — the JD emphasizes containers + Kubernetes. Spend 60% of time on Docker/Compose', 'Strategy 4: ')

doc.add_paragraph()
bold_body('Realistic outcome: ', 'You\'ll have hands-on experience with all 25 concepts and solid interview answers for each. You won\'t have muscle memory — that takes weeks. But you\'ll have enough to talk intelligently and demonstrate you\'ve done it.')

doc.add_page_break()

# ============================================================================
# DAY 1 — SATURDAY
# ============================================================================
doc.add_heading('DAY 1 (Saturday) — Linux Foundations [2 Hours]', level=1)
doc.add_paragraph()

# Block 1
doc.add_heading('Block 1: Users + Groups + Permissions (40 min)', level=2)
bold_body('Concepts covered: ', '#1 Filesystem, #2 Users, #3 Groups, #4 Permissions')
doc.add_paragraph()

bold_body('Do this (15 min):', '')
code('ls /etc /var/log /proc /opt /home                    # Know the hierarchy')
code('sudo useradd -m -s /bin/bash alice                   # Create user')
code('sudo useradd -r -s /usr/sbin/nologin svc-monitor     # Service account')
code('sudo groupadd soc-team && sudo usermod -aG soc-team alice')
code('sudo mkdir -p /opt/company/{app,logs}')
code('sudo chown root:soc-team /opt/company/logs')
code('sudo chmod 770 /opt/company/logs                     # Team-only access')
code('sudo chmod 2770 /opt/company/app                     # SGID for collaboration')
code('sudo -u alice touch /opt/company/logs/test.log       # Should WORK')
code('id alice && groups alice                             # Verify')

doc.add_paragraph()
bold_body('Memorize these answers (10 min):', '')
bullet('/etc = configs, /var/log = logs, /proc = kernel virtual fs, /opt = third-party apps', 'Filesystem: ')
bullet('useradd -m -s /bin/bash for humans, useradd -r -s /usr/sbin/nologin for services', 'Users: ')
bullet('usermod -aG (ALWAYS -a to append, never replace)', 'Critical flag: ')
bullet('770 = owner+group full, others nothing. SGID (2770) = new files inherit directory group. Sticky bit on /tmp = only owner deletes own files', 'Permissions: ')
bullet('Never chmod 777. Diagnose who needs access, use group ownership + 770.', '777 question: ')

doc.add_paragraph()
bold_body('Skip: ', 'The SUID deep dive, /etc/passwd field-by-field breakdown — read those on the bus/train.')

doc.add_paragraph()

# Block 2
doc.add_heading('Block 2: systemd + Cron + Logs (40 min)', level=2)
bold_body('Concepts covered: ', '#5 systemd, #6 Cron, #7 Log Analysis')
doc.add_paragraph()

bold_body('Do this (20 min):', '')
code('# Create a simple monitoring script')
code('sudo mkdir -p /opt/monitoring')
code('echo \'#!/bin/bash\nwhile true; do echo "$(date) OK"; sleep 60; done\' | sudo tee /opt/monitoring/health.sh')
code('sudo chmod +x /opt/monitoring/health.sh')
code('')
code('# Create systemd unit file')
code('sudo tee /etc/systemd/system/health-monitor.service << EOF')
code('[Unit]')
code('Description=Health Monitor')
code('After=network.target')
code('[Service]')
code('Type=simple')
code('ExecStart=/opt/monitoring/health.sh')
code('Restart=on-failure')
code('RestartSec=10')
code('[Install]')
code('WantedBy=multi-user.target')
code('EOF')
code('')
code('sudo systemctl daemon-reload')
code('sudo systemctl enable --now health-monitor')
code('sudo systemctl status health-monitor')
code('sudo journalctl -u health-monitor -n 5')
code('')
code('# Cron — add one job')
code('echo "0 2 * * * /opt/monitoring/backup.sh" | sudo crontab -')
code('sudo crontab -l')
code('')
code('# Log analysis — key patterns')
code('journalctl --since "1 hour ago" -p err')
code('journalctl -u ssh | grep "Failed password" | tail -5')

doc.add_paragraph()
bold_body('Memorize these answers (10 min):', '')
bullet('[Unit] = metadata + deps, [Service] = what to run + restart policy, [Install] = boot integration', 'Unit file: ')
bullet('daemon-reload after editing unit files, enable = boot, start = now, enable --now = both', 'Key commands: ')
bullet('0 2 * * * = daily at 2AM. Format: min hour day month weekday', 'Cron: ')
bullet('journalctl -u service for per-service, grep "Failed password" for brute-force, -p err for errors only', 'Logs: ')

doc.add_paragraph()

# Block 3
doc.add_heading('Block 3: Storage — fdisk + LVM + fstab (25 min)', level=2)
bold_body('Concepts covered: ', '#8 fdisk, #9 LVM, #10 fstab')
doc.add_paragraph()

bold_body('If you have an extra virtual disk, do this (15 min):', '')
code('# LVM is the critical skill — fdisk is just the prerequisite')
code('sudo pvcreate /dev/sdb')
code('sudo vgcreate vg_data /dev/sdb')
code('sudo lvcreate -L 2G -n lv_logs vg_data')
code('sudo mkfs.ext4 /dev/vg_data/lv_logs')
code('sudo mkdir -p /opt/logs')
code('sudo mount /dev/vg_data/lv_logs /opt/logs')
code('df -h /opt/logs')
code('')
code('# THE KEY SKILL — live expansion')
code('sudo lvextend -L +1G /dev/vg_data/lv_logs')
code('sudo resize2fs /dev/vg_data/lv_logs')
code('df -h /opt/logs    # Now 3GB, zero downtime')

doc.add_paragraph()
bold_body('If NO extra disk, just memorize (10 min):', '')
bullet('PV (raw disk) → VG (pool) → LV (flexible partition) → filesystem → mount', 'LVM hierarchy: ')
bullet('lvextend -L +5G /dev/vg/lv_name && resize2fs — done live, no unmount', '/var/log full fix: ')
bullet('Device, mount point, type, options (noexec, nosuid), dump, fsck order. Always backup before editing. Test with mount -a.', 'fstab: ')

doc.add_paragraph()

# Block 4
doc.add_heading('Block 4: iptables + SSH (15 min — READ ONLY)', level=2)
bold_body('Concepts covered: ', '#11-16 Firewall + SSH')
doc.add_paragraph()

bold_body('Don\'t type commands — just memorize answers:', '')
bullet('Default DROP on INPUT. Whitelist: loopback, ESTABLISHED/RELATED, SSH (22), HTTP (80/443). Log before implicit drop.', 'iptables: ')
bullet('-m recent tracks connections per IP. >3 new SSH connections in 60s = DROP + LOG.', 'Rate limiting: ')
bullet('ssh-keygen -t ed25519, ssh-copy-id to deploy, chmod 600 private key', 'SSH keys: ')
bullet('PasswordAuthentication no, PermitRootLogin no, MaxAuthTries 3', 'SSH hardening: ')
bullet('Keys can\'t be brute-forced or phished. Private key never crosses the network.', 'Why keys > passwords: ')

doc.add_page_break()

# ============================================================================
# DAY 2 — SUNDAY
# ============================================================================
doc.add_heading('DAY 2 (Sunday) — Docker + Compose + AWS Deployment [2 Hours]', level=1)
doc.add_paragraph()
bold_body('This is the high-priority day. ', 'The JD emphasizes containers, Kubernetes, and CI/CD. Spend more energy here.')

doc.add_paragraph()

# Block 1
doc.add_heading('Block 1: Dockerfile + Container Lifecycle (40 min)', level=2)
bold_body('Concepts covered: ', '#17 Dockerfile, #18 Lifecycle, #19 Volumes, #20 Networks')
doc.add_paragraph()

bold_body('Do this (25 min):', '')
code('mkdir -p ~/docker-lab && cd ~/docker-lab')
code('')
code('# Simple Flask app')
code('cat > app.py << \'PY\'')
code('from flask import Flask, jsonify')
code('import os')
code('app = Flask(__name__)')
code('@app.route("/health")')
code('def health(): return jsonify({"status": "ok", "host": os.uname().nodename})')
code('@app.route("/")')
code('def home(): return jsonify({"app": "AI-SOC Watchdog"})')
code('if __name__ == "__main__": app.run(host="0.0.0.0", port=5000)')
code('PY')
code('')
code('echo "flask==3.0.0\ngunicorn==21.2.0" > requirements.txt')
code('')
code('cat > Dockerfile << \'DF\'')
code('FROM python:3.11-slim')
code('WORKDIR /app')
code('COPY requirements.txt .')
code('RUN pip install --no-cache-dir -r requirements.txt')
code('COPY app.py .')
code('RUN useradd -r appuser')
code('USER appuser')
code('EXPOSE 5000')
code('ENTRYPOINT ["gunicorn", "--bind", "0.0.0.0:5000", "app:app"]')
code('DF')
code('')
code('docker build -t my-api:v1 .')
code('docker run -d --name api -p 5000:5000 my-api:v1')
code('curl http://localhost:5000/health')
code('')
code('# Lifecycle commands — try each one')
code('docker ps                          # Running containers')
code('docker logs api                    # App output')
code('docker exec api cat /etc/os-release  # Run command inside')
code('docker stop api && docker rm api')
code('')
code('# Volumes')
code('docker volume create app-data')
code('docker run -d --name api -p 5000:5000 -v app-data:/app/data my-api:v1')
code('')
code('# Networks')
code('docker network create app-net')
code('docker run -d --name api --network app-net my-api:v1')
code('docker run --rm --network app-net alpine ping -c 2 api  # DNS works!')

doc.add_paragraph()
bold_body('Memorize (10 min):', '')
bullet('FROM (base) → COPY requirements first (layer cache) → RUN pip install → COPY code → USER non-root → ENTRYPOINT', 'Dockerfile order: ')
bullet('ENTRYPOINT = fixed command. CMD = overridable default args.', 'ENTRYPOINT vs CMD: ')
bullet('Named volumes = Docker-managed, survive container removal. Bind mounts = host path mapped in.', 'Volumes: ')
bullet('Custom bridge gives DNS (containers find each other by name). Default bridge = IP only.', 'Networks: ')

doc.add_paragraph()

# Block 2
doc.add_heading('Block 2: Docker Compose (40 min)', level=2)
bold_body('Concepts covered: ', '#21-25 Compose YAML, management, networks, volumes, depends_on')
doc.add_paragraph()

bold_body('Do this (25 min):', '')
code('mkdir -p ~/compose-lab && cd ~/compose-lab')
code('mkdir -p app nginx')
code('')
code('# Copy app from Block 1')
code('cp ~/docker-lab/app.py ~/docker-lab/requirements.txt ~/docker-lab/Dockerfile app/')
code('')
code('# Nginx reverse proxy config')
code('cat > nginx/default.conf << \'CONF\'')
code('upstream api { server app:5000; }')
code('server {')
code('  listen 80;')
code('  location / { proxy_pass http://api; proxy_set_header Host $host; }')
code('}')
code('CONF')
code('')
code('# docker-compose.yml — THE KEY FILE')
code('cat > docker-compose.yml << \'YML\'')
code('version: "3.8"')
code('services:')
code('  nginx:')
code('    image: nginx:alpine')
code('    ports: ["80:80"]')
code('    volumes: ["./nginx/default.conf:/etc/nginx/conf.d/default.conf:ro"]')
code('    networks: [frontend]')
code('    depends_on: [app]')
code('    restart: unless-stopped')
code('  app:')
code('    build: ./app')
code('    networks: [frontend, backend]')
code('    depends_on: [redis]')
code('    restart: unless-stopped')
code('  redis:')
code('    image: redis:7-alpine')
code('    volumes: [redis-data:/data]')
code('    networks: [backend]')
code('    restart: unless-stopped')
code('networks:')
code('  frontend:')
code('  backend:')
code('    internal: true')
code('volumes:')
code('  redis-data:')
code('YML')
code('')
code('docker compose up -d --build')
code('docker compose ps')
code('curl http://localhost/health')
code('')
code('# Verify network isolation')
code('docker exec proxy ping -c 2 redis 2>&1   # FAILS — different networks')
code('docker exec app ping -c 2 redis           # WORKS — both on backend')
code('')
code('# Verify persistence')
code('docker compose down && docker compose up -d')
code('docker volume ls | grep redis              # Volume survived')
code('')
code('docker compose down')

doc.add_paragraph()
bold_body('Memorize (10 min):', '')
bullet('services (what to run), networks (how they connect), volumes (where data persists)', 'Compose structure: ')
bullet('internal: true = no external access. Containers on different networks can\'t talk.', 'Network isolation: ')
bullet('"down" keeps volumes. "down -v" DELETES volumes = data loss.', 'down vs down -v: ')
bullet('Controls startup ORDER only, not readiness. Use healthchecks for true readiness.', 'depends_on: ')
bullet('unless-stopped = restarts on crash but not after manual stop. always = always restarts.', 'Restart policies: ')

doc.add_paragraph()

# Block 3
doc.add_heading('Block 3: AWS Deployment for AI-SOC Watchdog (30 min READ)', level=2)
bold_body('This connects your project to the JD. ', 'You won\'t build EKS this weekend, but you need to EXPLAIN how you would.')
doc.add_paragraph()

bold_body('How the AI-SOC Watchdog maps to AWS + EKS:', '')
doc.add_paragraph()

bullet('Backend (Flask + Gunicorn) → Docker image → pushed to ECR → deployed as a Kubernetes Deployment on EKS', 'App Container: ')
bullet('Frontend (React build) → served via Nginx container or S3 + CloudFront', 'Frontend: ')
bullet('ChromaDB → runs as a sidecar container or separate pod with a PersistentVolumeClaim (EBS-backed)', 'Vector DB: ')
bullet('Supabase stays managed, OR migrate to RDS PostgreSQL for full AWS', 'Database: ')
bullet('S3 failover already built — just point to an S3 bucket in the same region', 'S3 Failover: ')

doc.add_paragraph()
bold_body('Terraform IaC — what you\'d write:', '')
bullet('VPC with public + private subnets across 2 AZs')
bullet('EKS cluster with managed node group (t3.medium, 2-3 nodes)')
bullet('ECR repository for the Docker image')
bullet('RDS PostgreSQL (if replacing Supabase)')
bullet('S3 bucket for failover storage')
bullet('ALB (Application Load Balancer) as ingress')
bullet('IAM roles for EKS nodes and pods (IRSA)')

doc.add_paragraph()
bold_body('Kubernetes manifests — what you\'d create:', '')
bullet('Deployment: 2 replicas of the Flask app, resource limits, health checks on /health', 'deployment.yaml: ')
bullet('ClusterIP service exposing port 5000 internally', 'service.yaml: ')
bullet('ALB ingress routing traffic to the service', 'ingress.yaml: ')
bullet('API keys, DB credentials stored as Kubernetes Secrets', 'secrets.yaml: ')
bullet('Non-sensitive config (ALLOWED_ORIGINS, LOG_LEVEL)', 'configmap.yaml: ')

doc.add_paragraph()
bold_body('CI/CD pipeline (GitHub Actions):', '')
bullet('on push to main → build Docker image → push to ECR → kubectl apply manifests → run health check')
bullet('Staging branch deploys to staging namespace, main deploys to production namespace')

doc.add_paragraph()
bold_body('Interview answer — "How would you deploy this to AWS?":', '')
p = doc.add_paragraph()
r = p.add_run('"I\'d containerize the backend with Docker (already done), push to ECR, and deploy to EKS with Terraform-managed infrastructure. The Terraform would provision a VPC, EKS cluster, and ALB. Kubernetes manifests define the deployment with 2 replicas, health checks on /health, and secrets for API keys. CI/CD via GitHub Actions: push to main triggers image build, ECR push, and kubectl rollout. ChromaDB runs as a sidecar with EBS-backed persistent storage. The frontend goes to S3 + CloudFront. The S3 failover system I already built slots right in."')
r.italic = True

doc.add_page_break()

# ============================================================================
# CHEAT SHEET
# ============================================================================
doc.add_heading('One-Page Cheat Sheet — Print This', level=1)

doc.add_heading('Linux', level=2)
code('useradd -m -s /bin/bash USER          # Create user')
code('useradd -r -s /usr/sbin/nologin SVC   # Service account')
code('usermod -aG GROUP USER                # Add to group (ALWAYS -a)')
code('chmod 770 DIR                         # Owner+group full access')
code('chmod 2770 DIR                        # SGID — inherit group')
code('systemctl enable --now SERVICE        # Start + boot persist')
code('systemctl status SERVICE              # Check status')
code('journalctl -u SERVICE -n 20           # Service logs')
code('journalctl -p err --since "1h ago"    # Recent errors')
code('pvcreate → vgcreate → lvcreate → mkfs → mount  # LVM flow')
code('lvextend -L +5G LV && resize2fs LV   # Live expand')
code('iptables -P INPUT DROP               # Default deny')
code('ssh-keygen -t ed25519                 # Generate key')

doc.add_heading('Docker', level=2)
code('docker build -t name:tag .            # Build image')
code('docker run -d --name X -p 80:80 IMG   # Run container')
code('docker ps / logs / exec / stop / rm   # Lifecycle')
code('docker volume create NAME             # Persistent storage')
code('docker network create NAME            # Custom network w/ DNS')

doc.add_heading('Docker Compose', level=2)
code('docker compose up -d --build          # Start stack')
code('docker compose ps / logs / down       # Manage')
code('docker compose down -v                # DESTRUCTIVE — deletes volumes')

doc.add_heading('AWS/K8s (Verbal)', level=2)
code('Terraform: VPC → EKS → ECR → ALB → RDS → S3')
code('K8s: Deployment (replicas) → Service → Ingress → Secrets')
code('CI/CD: push → build → ECR → kubectl apply → health check')

doc.add_page_break()

# ============================================================================
# SCHEDULE
# ============================================================================
doc.add_heading('Exact Schedule', level=1)

doc.add_heading('Saturday', level=2)
p = doc.add_paragraph()
r = p.add_run('Total: 2 hours')
r.bold = True

bullet('Users + Groups + Permissions — hands-on + memorize answers', '0:00-0:40 → ')
bullet('systemd + Cron + Logs — create one service, one cron job, grep logs', '0:40-1:20 → ')
bullet('LVM + fstab — create and expand a volume (or memorize if no extra disk)', '1:20-1:45 → ')
bullet('iptables + SSH — READ ONLY, memorize the interview answers', '1:45-2:00 → ')

doc.add_heading('Sunday', level=2)
p = doc.add_paragraph()
r = p.add_run('Total: 2 hours')
r.bold = True

bullet('Dockerfile + build + run + volumes + networks — build one image, test lifecycle', '0:00-0:40 → ')
bullet('Docker Compose — build the 3-service stack, verify network isolation', '0:40-1:20 → ')
bullet('AWS deployment plan — READ, memorize the interview answer', '1:20-1:50 → ')
bullet('Review: pick 5 random concepts, explain each in 30 seconds without notes', '1:50-2:00 → ')

doc.add_paragraph()
doc.add_paragraph()
green_text('You\'ve got this. 4 focused hours > 10 scattered hours.')

# Save
output_path = r'c:\Users\karan\Desktop\Weekend_Linux_Docker_AWS_Plan.docx'
doc.save(output_path)
print(f"Saved to: {output_path}")
